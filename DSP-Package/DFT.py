import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import *
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from signalcompare import SignalComapreAmplitude
from signalcompare import SignalComaprePhaseShift
import comparesignal2
# Enable interactive plotting
plt.ion()


def compare_result(input1, input2):
    file_path = "DFT-output-given/Output_Signal_DFT_A,Phase.txt"
    if file_path:
        with open(file_path, 'r') as file:
            lines = file.readlines()
            if len(lines) >= 3:
                array_size = int(lines[2].strip())
                amplitudes = []
                phases = []

                for line in lines[3:]:
                    values = line.strip().split()
                    if len(values) >= 2:
                        x_value = float(values[0])
                        y_value = float(values[1])
                        amplitudes.append(x_value)
                        phases.append(y_value)

        if SignalComapreAmplitude([round(x, 5) for x in input1], [round(x, 5) for x in amplitudes]) and SignalComaprePhaseShift(input2, phases):
            print("test passed successfully")


                    # Create a function to apply Fourier Transform to the input signal
def apply_fourier_transform(x_values, y_values, sampling_frequency):
    N = len(x_values)

    if sampling_frequency == 0:
        print("Sampling frequency cannot be zero.")
        return []

    T = 1 / sampling_frequency
    xk = []

    # Calculate the frequencies
    frequencies = [(2 * np.pi * k) / (N * T) for k in range(N)]

    for k in range(N):
        sum_xk = 0
        for n in range(N):
            sum_xk += y_values[n] * np.exp(-1j * 2 * np.pi * k * n / N)
        xk.append(sum_xk)

    return xk, frequencies

##remove dc task 5#########
def remove_dc_component(y_values):
    print("remove_dc_component function is called")
    mean_value = np.mean(y_values)
    y_values = [val - mean_value for val in y_values]
    dc_outputs = y_values.copy()
   # print(dc_inputs, "\n", dc_outputs)
    return y_values, dc_outputs

def remove_dc_component_FD(dft_result):
    # Set the first element (DC component) to zero
    dft_result[0] = 0
    return dft_result

def plot_amplitude_phase_in_DC(x_values, y_values, sampling_frequency, title=None):
    xk, frequencies = apply_fourier_transform(x_values, y_values, sampling_frequency)
    xk= remove_dc_component_FD(xk)
    amplitudes = [np.abs(xi) for xi in xk]
    phases = [np.angle(xi) for xi in xk]
    # Create subplots with some space between them
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6))
    fig.subplots_adjust(hspace=0.5)  # Adjust vertical spacing

    # Plot amplitude as rectangles in the first subplot
    ax1.bar(frequencies, amplitudes, width=0.2, color='orange', label='Amplitude')
    ax1.set_title("Frequency vs Amplitude")
    ax1.set_xlabel("Frequency (Hz)")
    ax1.set_ylabel("Amplitude")
    ax1.grid()
    ax1.legend()

    # Plot phase as rectangles in the second subplot
    ax2.bar(frequencies, phases, width=0.2, color='green', label='Phase')
    ax2.set_title("Frequency vs Phase")
    ax2.set_xlabel("Frequency (Hz)")
    ax2.set_ylabel("Phase (radians)")
    ax2.grid()
    ax2.legend()

    if title:
        fig.suptitle(title, fontsize=16)  # Add title to the entire figure

    plt.show()

    return frequencies, amplitudes, phases
    pass

# Create a function to plot amplitude and phase as rectangles
def plot_amplitude_phase(x_values, y_values, sampling_frequency, title=None):
    xk, frequencies = apply_fourier_transform(x_values, y_values, sampling_frequency)
    amplitudes = [np.abs(xi) for xi in xk]
    phases = [np.angle(xi) for xi in xk]

    # Create subplots with some space between them
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6))
    fig.subplots_adjust(hspace=0.5)  # Adjust vertical spacing

    # Plot amplitude as rectangles in the first subplot
    ax1.bar(frequencies, amplitudes, width=0.2, color='orange', label='Amplitude')
    ax1.set_title("Frequency vs Amplitude")
    ax1.set_xlabel("Frequency (Hz)")
    ax1.set_ylabel("Amplitude")
    ax1.grid()
    ax1.legend()

    # Plot phase as rectangles in the second subplot
    ax2.bar(frequencies, phases, width=0.2, color='green', label='Phase')
    ax2.set_title("Frequency vs Phase")
    ax2.set_xlabel("Frequency (Hz)")
    ax2.set_ylabel("Phase (radians)")
    ax2.grid()
    ax2.legend()

    if title:
        fig.suptitle(title, fontsize=16)  # Add title to the entire figure

    plt.show()



    return frequencies, amplitudes, phases




def plot_DC_signal(x_values, y_values, dc_outputs, title):
    fig, ax = plt.subplots(figsize=(12, 6))  # Adjust the figure size
    ax.plot(x_values, y_values, label='Original Signal')
    ax.plot(x_values, dc_outputs, label='Signal After DC Removal')
    ax.axhline(0, color='black', linewidth=0.8, linestyle='--')  # Add a horizontal line at the x-axis
    ax.set_title(title)
    ax.set_xlabel('Time')
    ax.set_ylabel('Amplitude')
    ax.legend()
    plt.show()


# Create a function to open the file, process the signal, and generate the Word document
def open_file():
    file_path = filedialog.askopenfilename()
    option=combo_box.get()
    if file_path:
        with open(file_path, 'r') as file:
            lines = file.readlines()
            if len(lines) >= 3:
                array_size = int(lines[2].strip())
                x_values = []  # Initialize the x_values array
                y_values = []  # Initialize the y_values array

                for line in lines[3:]:  # Iterate from the fourth line
                    values = line.strip().split()
                    if len(values) >= 2:  # Ensure there are at least 2 values on the line
                        x_value = float(values[0])  # Convert to float
                        y_value = float(values[1])  # Convert to float
                        x_values.append(x_value)  # Append the first value to x_values
                        y_values.append(y_value)  # Append the second value to y_values

                # Retrieve the sampling frequency from the Entry widget
                sampling_frequency = float(sampling_frequency_entry.get())

                if(option=="DC Remove"):
                    dc_inputs, dc_outputs = remove_dc_component(y_values)
                    rounded_dc_outputs = [round(val, 4) for val in dc_outputs]
                    plot_DC_signal(x_values,y_values, dc_outputs, 'Signal Before-After DC Removal')
                    ##plot frequency domain before delete dc component
                    plot_amplitude_phase(x_values, y_values, sampling_frequency,title=" frequency domain before delete DC component")
                    ##plot frequency domain after delete dc component
                    plot_amplitude_phase_in_DC(x_values,rounded_dc_outputs,sampling_frequency,title=" frequency domain after delete DC component")
                    comparesignal2.SignalSamplesAreEqual("input_output_files/DC_component_output.txt",rounded_dc_outputs)

                elif(option=="DFT"):
                    frequencies, amplitudes, phases = plot_amplitude_phase(x_values,y_values, sampling_frequency)
                    compare_result(amplitudes,phases)

                    # Create a Word document and add amplitude and phase data
                    document = Document()
                    document.add_heading('Amplitude and Phase Data', 0)

                    for freq, amp, phase in zip(frequencies, amplitudes, phases):
                        document.add_paragraph(f'Frequency (Hz): {freq:.4f}\tAmplitude: {amp:.13f}\tPhase (radians): {phase:.13f}')

                    # Save the Word document
                    document.save("amplitude_phase_data.docx")

# Create the GUI
root = tk.Tk()
root.geometry("500x400")
root.configure(bg="beige")
root.title("Fourier Transform of a Signal")



sampling_frequency_label = tk.Label(root, text="Enter the sampling frequency (Hz):")
sampling_frequency_label.configure(font=("Arial", 12), bg="beige")
sampling_frequency_label.place(relx=0.5, rely=0.2, anchor=tk.CENTER)

sampling_frequency_entry = tk.Entry(root)
sampling_frequency_entry.configure(font=("Arial", 12))
sampling_frequency_entry.place(relx=0.5, rely=0.3, anchor=tk.CENTER)

signals_operations = ["DFT", "DC Remove"]
style = ttk.Style()
style.configure("TCombobox", background="black", foreground="darkred")
combo_box = ttk.Combobox(root, values=signals_operations, style="TCombobox")
combo_box.place(relx=0.34, rely=0.35)

# Create and place the button
browse_and_process_button = tk.Button(root, text="Apply Operation", command=open_file)
browse_and_process_button.configure(font=("Arial", 12), bg="black", fg="white")
browse_and_process_button.place(relx=0.47, rely=0.55, anchor=tk.CENTER)
root.mainloop()
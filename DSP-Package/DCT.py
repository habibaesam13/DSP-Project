import tkinter as tk
from tkinter import filedialog
import numpy as np
import matplotlib.pyplot as plt
from  comparesignal2 import SignalSamplesAreEqual
import matplotlib.pyplot as plt
from docx import Document

# Enable interactive plotting
plt.ion()
dct_result = 0





def open_file():
    global dct_result  # Make dct_result a global variable to store the result
    file_path = filedialog.askopenfilename()
    if file_path:
        with open(file_path, 'r') as file:
            lines = file.readlines()
            if len(lines) >= 3:
                array_size = int(lines[2].strip())
                x_values = []  # Initialize the x_values array
                y_values = []  # Initialize the y_values array

                for line in lines[3:]:  # Iterate from the fourth line
                    values = line.strip().split()
                    if len(values) >= 2:  # Ensure there are at least 2 values on the line
                        x_value = float(values[0])  # Convert to float
                        y_value = float(values[1])  # Convert to float
                        x_values.append(x_value)  # Append the first value to x_values
                        y_values.append(y_value)  # Append the second value to y_values

                # Check if the number of samples matches the expected size
                if len(x_values) == array_size and len(y_values) == array_size:
                    # Apply 1D DCT to y_values
                    dct_result = dct1d(y_values)
                    print("Original Samples:", y_values)
                    print("DCT Coefficients:", dct_result)

                    document = Document()
                    document.add_heading('DCT Results', level=1)

                    # Add information about original samples
                    document.add_heading('Original Samples:', level=2)
                    document.add_paragraph(str(y_values))

                    # Add information about DCT coefficients
                    document.add_heading('DCT Coefficients:', level=2)
                    document.add_paragraph(str(dct_result))

                    # Save the Word document
                    document.save("DCT.docx")




                    SignalSamplesAreEqual("input_output_files/DCT_output.txt",dct_result)

                    # Plot results
                    plot_results(x_values, y_values, dct_result)
                else:
                    print("The number of samples in the file does not match the expected size.")


def dct1d(x):
    N = len(x)
    y = np.zeros(N)

    for k in range(N):
        sum_val = 0.0
        for n in range(N):
            sum_val += np.sqrt(2/N) * x[n] * np.cos(np.pi / (4 * N) * (2 * n - 1) * (2 * k - 1))
        y[k] = sum_val


    return y



def plot_results(x_values, y_values, dct_result):
    plt.figure(figsize=(10, 6))

    # Plot original samples
    plt.subplot(3, 1, 1)
    plt.plot(x_values, y_values, label='Original Samples')
    plt.title('Original Samples')
    plt.xlabel('Index')
    plt.ylabel('Value')
    plt.legend()

    # Plot DCT Coefficients
    plt.subplot(3, 1, 3)
    plt.plot(dct_result, label='DCT Coefficients')
    plt.title('DCT Coefficients')
    plt.xlabel('DCT Index')
    plt.ylabel('Value')
    plt.legend()





# GUI
root = tk.Tk()
root.geometry("500x400")
root.configure(bg="beige")
root.title("Discrete Cosine Transform (DCT)")

upload_button_s1 = tk.Button(root, text="Upload SIGNAL", command=open_file)
upload_button_s1.configure(font=("Arial", 12), bg="black", fg="white")
upload_button_s1.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

root.mainloop()

import tkinter as tk
from tkinter import messagebox, filedialog
import numpy as np
from practical.Practical_task1.CompareSignal import Compare_Signals
import matplotlib.pyplot as plt
import math
from docx import Document


Filter_type = ""
filter_widgets = []
x_values = []
y_values = []
Hd = []
signal_after_filtering = []
convRes = []


def test_filter_coff(filterType, filter_indices, coff, signalAfterFilterIndices, filteredSignalOutput):
    if filterType == "lowPass":
        print("\nTEST LOW PASS FILTER COEFFICIENTS\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 1\LPFCoefficients.txt",
            filter_indices, coff)
        print("\nTEST SIGNAL WITH LOW PASS FILTER\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 2\ecg_low_pass_filtered.txt",
            signalAfterFilterIndices, filteredSignalOutput)
    elif filterType == "highPass":
        print("\nTEST HIGH PASS FILTER COEFFICIENTS\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 3\HPFCoefficients.txt",
            filter_indices, coff)
        print("\nTEST SIGNAL WITH HIGH PASS FILTER\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 4\ecg_high_pass_filtered.txt",
            signalAfterFilterIndices, filteredSignalOutput)
    elif filterType == "bandPass":
        print("\nTEST BAND PASS FILTER COEFFICIENTS\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 5\BPFCoefficients.txt",
            filter_indices, coff)
        print("\nTEST SIGNAL WITH BAND PASS FILTER\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 6\ecg_band_pass_filtered.txt",
            signalAfterFilterIndices, filteredSignalOutput)
    elif filterType == "bandReject":
        print("\nTEST BAND REJECT FILTER COEFFICIENTS\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 7\BSFCoefficients.txt",
            filter_indices, coff)
        print("\nTEST SIGNAL WITH BAND REJECT FILTER\n")
        Compare_Signals(
            "D:\python\DSP-Tasks\DSP-Package\practical\Practical_task1\FIR test cases\Testcase 8\ecg_band_stop_filtered.txt",
            signalAfterFilterIndices, filteredSignalOutput)


def convolution(x_values, y_values, x_values2, y_values2):
    len1 = len(y_values)
    len2 = len(y_values2)

    start_index = int(min(x_values) + min(x_values2))
    end_index = int(max(x_values) + max(x_values2))

    indices = list(range(start_index, end_index + 1))

    for n in range(len1 + len2 - 1):
        sum_val = 0
        for m in range(min(n, len1 - 1) + 1):
            if 0 <= n - m < len2:
                sum_val += y_values[m] * y_values2[n - m]
        convRes.append(round(sum_val, 10))

    return indices, convRes


def apply_window(Hd, windowName, N):
    positiveindices = []
    multi_res = []
    if windowName == "Rectangular":
        # Return indices and values symmetrically
        indices = list(range(-int(N / 2), int(N / 2) + 1))
        return indices, Hd + Hd[::-1]
    elif windowName == "Hanning":
        for i in range(int(N / 2) + 1):
            positiveindices.append(0.5 + 0.5 * math.cos(2 * math.pi * i / N))
    elif windowName == "Hamming":
        for i in range(int(N / 2) + 1):
            positiveindices.append(0.54 + 0.46 * math.cos(2 * math.pi * i / N))
    elif windowName == "Blackman":
        for i in range(int(N / 2) + 1):
            component1 = 0.42
            component2 = 0.5 * math.cos(2 * math.pi * i / (N - 1))
            component3 = 0.08 * math.cos(4 * math.pi * i / (N - 1))

            positiveindices.append(component1 + component2 + component3)
    else:
        raise ValueError("Unsupported windowName")

    # Multiply corresponding Hd and window values
    for value1, value2 in zip(Hd, positiveindices):
        multi_res.append(round(value1 * value2, 10))

    # Iterate over negative indices and append to the list
    for i in range(-1, -len(multi_res) - 1, -1):
        value = multi_res[i]
        signal_after_filtering.append(value)

    # Iterate over positive indices and append to the list
    for i in range(1, len(multi_res)):
        value = multi_res[i]
        signal_after_filtering.append(value)

    # Return the indices and symmetric values
    indices = list(range(-int(N / 2), int(N / 2) + 1))

    return indices, signal_after_filtering


def prepare_filter():
    samplingFreq = float(sampling_frequency_entry.get())
    stopAttenuation = float(stop_attenuation_entry.get())
    transitionBand = float(transition_band_entry.get())
    # normalize
    transitionBand /= samplingFreq
    # stopBandAttenuation normalizedTransitionWidth WindowName
    window_stopBandAttenuation = [
        (21, 0.9, "Rectangular"),
        (44, 3.1, "Hanning"),
        (53, 3.3, "Hamming"),
        (74, 5.5, "Blackman"),
    ]
    for tup in window_stopBandAttenuation:
        first_value = tup[0]
        if first_value >= stopAttenuation:
            windowName = tup[2]
            N = round(tup[1] / transitionBand)
            if N % 2 == 0:
                N += 1
            break


    if Filter_type == "lowPass" or Filter_type == "highPass":
        cutoffFrequency = float(filter_widgets[1].get())
        if Filter_type == "lowPass":
            cutoffFrequency = (cutoffFrequency / samplingFreq) + (0.5 * transitionBand)
            Hd.append(2 * cutoffFrequency)
            for i in range(1, int(N / 2) + 1):
                Hd.append(
                    2 * cutoffFrequency * (math.sin(i * 2 * math.pi * cutoffFrequency) / (i * 2 * math.pi * cutoffFrequency)))
            indices, signal_after_filtering = apply_window(Hd, windowName, N)

        elif Filter_type == "highPass":
            # fc`=fc-0.5 transition band
            cutoffFrequency = (cutoffFrequency / samplingFreq) - (0.5 * transitionBand)
            Hd.append(1 - 2 * cutoffFrequency)
            for i in range(1, int(N / 2) + 1):
                Hd.append(
                    -2 * cutoffFrequency * (math.sin(i * 2 * math.pi * cutoffFrequency) / (i * 2 * math.pi * cutoffFrequency)))
            indices, signal_after_filtering = apply_window(Hd, windowName, N)

    elif Filter_type == "bandPass" or Filter_type == "bandReject":
        frequency1 = float(filter_widgets[1].get())  # Get frequency 1
        frequency2 = float(filter_widgets[3].get())  # Get frequency 2
        frequency1 /= samplingFreq
        frequency2 /= samplingFreq
        if Filter_type == "bandPass":
            #high pass     low pass
            frequency1 -= (0.5 * transitionBand)
            frequency2 += (0.5 * transitionBand)
            Hd.append(2 * (frequency2 - frequency1))
            for i in range(1, int(N / 2) + 1):
                inserted_val= 2 * frequency2 * (math.sin(i * 2 * math.pi * frequency2) / (i * 2 * math.pi * frequency2))
                inserted_val+= -2 * frequency1 * (math.sin(i * 2 * math.pi * frequency1) / (i * 2 * math.pi * frequency1))
                Hd.append(inserted_val)
            indices, signal_after_filtering = apply_window(Hd, windowName, N)
        elif Filter_type == "bandReject":
            #low pass      high pass
            frequency1 += (0.5 * transitionBand)
            frequency2 -= (0.5 * transitionBand)
            Hd.append(1 - 2 * (frequency2 - frequency1))
            for i in range(1, int(N / 2) + 1):
                inserted_val = 2 * frequency1 * (math.sin(i * 2 * math.pi * frequency1) / (i * 2 * math.pi * frequency1))
                inserted_val += -2 * frequency2 * (math.sin(i * 2 * math.pi * frequency2) / (i * 2 * math.pi * frequency2))
                Hd.append(inserted_val)
            indices, signal_after_filtering = apply_window(Hd, windowName, N)

    last_indices, lastRes = convolution(x_values, y_values, indices, signal_after_filtering)


    test_filter_coff(Filter_type, indices, signal_after_filtering, last_indices, lastRes)

    # # Plot the signals
    plot_signals(x_values, y_values, last_indices, lastRes)

    # Create a Word document and add amplitude and phase data
    document = Document()
    document.add_heading('Signal after Filtering ', 0)

    for Index, convolvedSignal in zip(last_indices, lastRes):
        document.add_paragraph(f'Index: {Index:.4f}\tFiltered value: {convolvedSignal:.10f}\t')

    # Save the Word document
    document.save(f"Convolved Signal With {Filter_type } and window {windowName}.docx")


# Create a function to open the file
def open_file():
    if not (lowPass_var.get() or highPass_var.get() or bandPass_var.get() or bandReject_var.get()):
        messagebox.showinfo("Alert", "Choose Filter Type")
        return

    file_path = filedialog.askopenfilename()

    if file_path:
        with open(file_path, "r") as file:
            lines = file.readlines()
            if len(lines) >= 3:
                array_size = int(lines[2].strip())

                for line in lines[3:]:
                    values = line.strip().split()
                    if len(values) >= 2:
                        x_value = float(values[0])
                        y_value = float(values[1])
                        x_values.append(x_value)
                        y_values.append(y_value)

    prepare_filter()


def on_checkbox_clicked():
    global Filter_type
    clear_filter_widgets()  # Clear old filter widgets

    if lowPass_var.get():
        Filter_type = "lowPass"
        uncheck_other_checkboxes(lowPass_checkbox)
        show_filter_widgets("Cut Off Frequency")

    if highPass_var.get():
        Filter_type = "highPass"
        uncheck_other_checkboxes(highPass_checkbox)
        show_filter_widgets("Cut Off Frequency")

    if bandPass_var.get():
        Filter_type = "bandPass"
        uncheck_other_checkboxes(bandPass_checkbox)
        show_filter_widgets("Frequency 1", "Frequency 2")

    if bandReject_var.get():
        Filter_type = "bandReject"
        uncheck_other_checkboxes(bandReject_checkbox)
        show_filter_widgets("Frequency 1", "Frequency 2")


def create_filter_widgets(*labels):
    global filter_widgets
    row_height = 0.55
    for label_text in labels:
        label = tk.Label(root, text=label_text + ":")
        label.configure(font=("Arial", 12, "bold"), fg="#053B50", bg="#EBE3D5")
        label.place(relx=0.18, rely=row_height, anchor=tk.CENTER)

        entry = tk.Entry(root)
        entry.configure(font=("Arial", 12, "bold"))
        entry.place(relx=0.55, rely=row_height, anchor=tk.CENTER)

        filter_widgets.append(label)
        filter_widgets.append(entry)

        row_height += 0.1


def clear_filter_widgets():
    global filter_widgets
    for widget in filter_widgets:
        widget.place_forget()  # Hide the widget
    filter_widgets = []  # Clear the list


def show_filter_widgets(*labels):
    create_filter_widgets(*labels)


def uncheck_other_checkboxes(clicked_checkbox):
    checkboxes = [lowPass_checkbox, highPass_checkbox, bandPass_checkbox, bandReject_checkbox]
    for checkbox in checkboxes:
        if checkbox != clicked_checkbox:
            checkbox.deselect()


def plot_signals(x_values, y_values, filtered_x, filtered_y):
    plt.figure(figsize=(10, 6))

    # Plot the original signal
    plt.subplot(2, 1, 1)
    plt.plot(x_values, y_values, label='Original Signal', color='blue')
    plt.title('Original Signal')
    plt.xlabel('Time')
    plt.ylabel('Amplitude')
    plt.legend()

    # Plot the filtered signal
    plt.subplot(2, 1, 2)
    plt.plot(filtered_x, filtered_y, label='Filtered Signal', color='red')
    plt.title('Filtered Signal')
    plt.xlabel('Time')
    plt.ylabel('Amplitude')
    plt.legend()

    plt.tight_layout()
    plt.show()


############################## Create the GUI  #########################
root = tk.Tk()
root.geometry("700x500")
root.configure(bg="#EBE3D5")
root.title("FILTERING")

# Sampling Frequency Entry and Label
sampling_frequency_label = tk.Label(root, text="Sampling Frequency(HZ):")
sampling_frequency_label.configure(font=("Arial", 12, "bold"), fg="#053B50", bg="#EBE3D5")
sampling_frequency_label.place(relx=0.194, rely=0.1, anchor=tk.CENTER)
sampling_frequency_entry = tk.Entry(root)
sampling_frequency_entry.configure(font=("Arial", 12, "bold"))
sampling_frequency_entry.place(relx=0.6, rely=0.1, anchor=tk.CENTER)

# Stop Attenuation Entry and Label
stop_attenuation_label = tk.Label(root, text="Stop Attenuation:")
stop_attenuation_label.configure(font=("Arial", 12, "bold"), fg="#053B50", bg="#EBE3D5")
stop_attenuation_label.place(relx=0.15, rely=0.18, anchor=tk.CENTER)
stop_attenuation_entry = tk.Entry(root)
stop_attenuation_entry.configure(font=("Arial", 12, "bold"))
stop_attenuation_entry.place(relx=0.6, rely=0.18, anchor=tk.CENTER)

# Transition Band Entry and Label
transition_band_label = tk.Label(root, text="Transition Band:")
transition_band_label.configure(font=("Arial", 12, "bold"), fg="#053B50", bg="#EBE3D5")
transition_band_label.place(relx=0.15, rely=0.26, anchor=tk.CENTER)
transition_band_entry = tk.Entry(root)
transition_band_entry.configure(font=("Arial", 12, "bold"))
transition_band_entry.place(relx=0.6, rely=0.26, anchor=tk.CENTER)

# Choose filter type
lowPass_var = tk.BooleanVar()
highPass_var = tk.BooleanVar()
bandPass_var = tk.BooleanVar()
bandReject_var = tk.BooleanVar()

# Checkbuttons for filter types
lowPass_checkbox = tk.Checkbutton(root, text="LOW PASS FILTER", variable=lowPass_var, command=on_checkbox_clicked)
highPass_checkbox = tk.Checkbutton(
    root, text="HIGH PASS FILTER", variable=highPass_var, command=on_checkbox_clicked
)
bandPass_checkbox = tk.Checkbutton(
    root, text="BAND PASS FILTER", variable=bandPass_var, command=on_checkbox_clicked
)
bandReject_checkbox = tk.Checkbutton(
    root, text="BAND REJECT FILTER", variable=bandReject_var, command=on_checkbox_clicked
)

# Place the filter type checkboxes in the window
lowPass_checkbox.configure(font=("Arial", 11, "bold"), fg="#053B50", bg="#EBE3D5")
lowPass_checkbox.place(relx=0.055, rely=0.38)
highPass_checkbox.configure(font=("Arial", 11, "bold"), fg="#053B50", bg="#EBE3D5")
highPass_checkbox.place(relx=0.4, rely=0.38)
bandPass_checkbox.configure(font=("Arial", 11, "bold"), fg="#053B50", bg="#EBE3D5")
bandPass_checkbox.place(relx=0.055, rely=0.45)
bandReject_checkbox.configure(font=("Arial", 11, "bold"), fg="#053B50", bg="#EBE3D5")
bandReject_checkbox.place(relx=0.4, rely=0.45)

# Create and place the button
browse_and_button = tk.Button(root, text="Upload Signal", command=open_file)
browse_and_button.configure(font=("Arial", 12, "bold"), bg="#EBE3D5", fg="#053B50")
browse_and_button.place(relx=0.45, rely=0.75, anchor=tk.CENTER)

root.mainloop()
